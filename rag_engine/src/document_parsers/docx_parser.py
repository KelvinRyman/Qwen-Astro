"""
DOCX文档解析器

支持Microsoft Word DOCX格式文档的文本提取、元数据提取和智能分块。
"""

import logging
from typing import List, Optional
from pathlib import Path
from .base_parser import DocumentParser, DocumentMetadata, DocumentChunk

try:
    from docx import Document
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError as e:
    logging.warning(f"DOCX解析库导入失败: {e}")
    DOCX_AVAILABLE = False


class DocxParser(DocumentParser):
    """DOCX文档解析器"""
    
    def __init__(self, config=None):
        super().__init__(config)
        self.extract_headers = self.config.get('extract_headers', True)
        self.extract_tables = self.config.get('extract_tables', True)
        self.min_paragraph_length = self.config.get('min_paragraph_length', 10)
        self.chunk_by_heading = self.config.get('chunk_by_heading', True)
    
    def can_parse(self, file_path: str) -> bool:
        """检查是否能解析DOCX文件"""
        if not DOCX_AVAILABLE:
            return False
        
        return Path(file_path).suffix.lower() == '.docx'
    
    def get_supported_extensions(self) -> List[str]:
        """获取支持的文件扩展名"""
        return ['.docx']
    
    def extract_metadata(self, file_path: str) -> DocumentMetadata:
        """提取DOCX元数据"""
        metadata = DocumentMetadata(format_type='docx')
        
        try:
            doc = Document(file_path)
            
            # 核心属性
            if hasattr(doc, 'core_properties'):
                core_props = doc.core_properties
                metadata.title = core_props.title
                metadata.author = core_props.author
                metadata.subject = core_props.subject
                metadata.keywords = core_props.keywords
                metadata.creator = core_props.creator
                
                # 日期
                if core_props.created:
                    metadata.creation_date = core_props.created.strftime('%Y-%m-%d %H:%M:%S')
                if core_props.modified:
                    metadata.modification_date = core_props.modified.strftime('%Y-%m-%d %H:%M:%S')
            
            # 统计信息
            word_count = 0
            paragraph_count = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraph_count += 1
                    word_count += len(paragraph.text.split())
            
            metadata.word_count = word_count
            metadata.extra_metadata = {
                'paragraph_count': paragraph_count,
                'table_count': len(doc.tables) if self.extract_tables else 0
            }
            
            # 文件大小
            metadata.file_size = Path(file_path).stat().st_size
            
        except Exception as e:
            self.logger.error(f"提取DOCX元数据失败: {e}")
        
        return metadata
    
    def extract_text(self, file_path: str) -> str:
        """提取DOCX全文"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text and len(text) >= self.min_paragraph_length:
                    text_parts.append(text)
            
            # 提取表格文本
            if self.extract_tables:
                for table in doc.tables:
                    table_text = self._extract_table_text(table)
                    if table_text:
                        text_parts.append(table_text)
            
            return self._clean_text('\n\n'.join(text_parts))
            
        except Exception as e:
            self.logger.error(f"提取DOCX文本失败: {e}")
            return ""
    
    def extract_chunks(self, file_path: str) -> List[DocumentChunk]:
        """提取DOCX分块"""
        chunks = []
        base_metadata = {'file_name': Path(file_path).name, 'format_type': 'docx'}
        
        try:
            doc = Document(file_path)
            
            if self.chunk_by_heading:
                chunks = self._extract_chunks_by_heading(doc, base_metadata)
            else:
                chunks = self._extract_chunks_by_paragraph(doc, base_metadata)
            
        except Exception as e:
            self.logger.error(f"提取DOCX分块失败: {e}")
        
        return chunks
    
    def _extract_chunks_by_heading(self, doc, base_metadata: dict) -> List[DocumentChunk]:
        """按标题层级分块"""
        chunks = []
        current_section = {'title': None, 'content': [], 'level': 0}
        section_counter = 0
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            # 检查是否为标题
            heading_level = self._get_heading_level(paragraph)
            
            if heading_level > 0:
                # 保存当前段落
                if current_section['content']:
                    chunk = self._create_section_chunk(
                        current_section, base_metadata, section_counter
                    )
                    if chunk:
                        chunks.append(chunk)
                    section_counter += 1
                
                # 开始新段落
                current_section = {
                    'title': text,
                    'content': [],
                    'level': heading_level
                }
            else:
                # 添加到当前段落
                if len(text) >= self.min_paragraph_length:
                    current_section['content'].append(text)
        
        # 处理最后一个段落
        if current_section['content']:
            chunk = self._create_section_chunk(
                current_section, base_metadata, section_counter
            )
            if chunk:
                chunks.append(chunk)
        
        return chunks
    
    def _extract_chunks_by_paragraph(self, doc, base_metadata: dict) -> List[DocumentChunk]:
        """按段落分块"""
        chunks = []
        paragraph_counter = 0
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text and len(text) >= self.min_paragraph_length:
                cleaned_text = self._clean_text(text)
                
                chunk_metadata = self._create_chunk_metadata(
                    base_metadata,
                    {
                        'paragraph_number': paragraph_counter + 1,
                        'chunk_type': 'paragraph'
                    }
                )
                
                chunk = DocumentChunk(
                    text=cleaned_text,
                    metadata=chunk_metadata,
                    chunk_id=f"paragraph_{paragraph_counter + 1}",
                    chunk_type='paragraph'
                )
                chunks.append(chunk)
                paragraph_counter += 1
        
        # 处理表格
        if self.extract_tables:
            for table_idx, table in enumerate(doc.tables):
                table_text = self._extract_table_text(table)
                if table_text:
                    cleaned_text = self._clean_text(table_text)
                    
                    chunk_metadata = self._create_chunk_metadata(
                        base_metadata,
                        {
                            'table_number': table_idx + 1,
                            'chunk_type': 'table'
                        }
                    )
                    
                    chunk = DocumentChunk(
                        text=cleaned_text,
                        metadata=chunk_metadata,
                        chunk_id=f"table_{table_idx + 1}",
                        chunk_type='table'
                    )
                    chunks.append(chunk)
        
        return chunks
    
    def _get_heading_level(self, paragraph) -> int:
        """获取段落的标题级别"""
        try:
            style_name = paragraph.style.name.lower()
            if 'heading' in style_name:
                # 提取数字
                import re
                match = re.search(r'heading\s*(\d+)', style_name)
                if match:
                    return int(match.group(1))
                return 1  # 默认为1级标题
        except Exception:
            pass
        
        return 0  # 不是标题
    
    def _create_section_chunk(self, section: dict, base_metadata: dict, section_id: int) -> Optional[DocumentChunk]:
        """创建段落分块"""
        if not section['content']:
            return None
        
        # 组合标题和内容
        text_parts = []
        if section['title']:
            text_parts.append(section['title'])
        text_parts.extend(section['content'])
        
        cleaned_text = self._clean_text('\n\n'.join(text_parts))
        
        chunk_metadata = self._create_chunk_metadata(
            base_metadata,
            {
                'section_title': section['title'],
                'section_level': section['level'],
                'section_number': section_id + 1,
                'chunk_type': 'section'
            }
        )
        
        return DocumentChunk(
            text=cleaned_text,
            metadata=chunk_metadata,
            chunk_id=f"section_{section_id + 1}",
            section_title=section['title'],
            chunk_type='section'
        )
    
    def _extract_table_text(self, table) -> str:
        """提取表格文本"""
        try:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_data.append(cell_text)
                if row_data:
                    table_data.append(' | '.join(row_data))
            
            return '\n'.join(table_data)
        except Exception as e:
            self.logger.warning(f"提取表格文本失败: {e}")
            return ""
